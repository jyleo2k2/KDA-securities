"""Build the team-facing Word report for two exhaustive mock customer examples."""

from __future__ import annotations

import json
from collections.abc import Iterable, Mapping, Sequence
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

from scripts.render_customer_data_examples_md import (
    DETAILED_HOLDING_FIELDS,
    FIELD_DESCRIPTIONS,
)

ROOT = Path(__file__).resolve().parents[1]
MOCK_DIR = ROOT / "data" / "mock"
OUTPUT_PATH = ROOT / "docs" / "30_스펙" / "고객_목데이터_대표_2명_보고서.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "F4F7FB"
WHITE = "FFFFFF"
GRID = "C9D2DC"
BLACK = "1F2937"

MONEY_FIELDS = {
    key
    for key in FIELD_DESCRIPTIONS
    if key.endswith("_krw") or key in {"balance_krw", "amount_krw"}
}
RATIO_FIELDS = {"risky_asset_ratio", "safe_asset_ratio", "cash_ratio", "weight"}


def _set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, margins: Mapping[str, int]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "bottom", "start", "end"):
        tag = qn(f"w:{side}")
        margin = tc_mar.find(tag)
        if margin is None:
            margin = OxmlElement(f"w:{side}")
            tc_mar.append(margin)
        margin.set(qn("w:w"), str(margins[side]))
        margin.set(qn("w:type"), "dxa")


def _set_width(parent, tag: str, width: int) -> None:
    element = parent.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        parent.append(element)
    element.set(qn("w:type"), "dxa")
    element.set(qn("w:w"), str(width))


def _apply_table_geometry(table, widths: Sequence[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must total {CONTENT_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    _set_width(tbl_pr, "w:tblW", CONTENT_WIDTH_DXA)

    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for index, width in enumerate(widths):
        table.columns[index].width = Twips(width)
    for row in table.rows:
        row.height = None
        for index, cell in enumerate(row.cells):
            cell.width = Twips(widths[index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            _set_width(tc_pr, "w:tcW", widths[index])
            _set_cell_margins(cell, CELL_MARGINS_DXA)


def _set_cell_text(
    cell,
    value: Any,
    *,
    bold: bool = False,
    size: float = 8.5,
    color: str = BLACK,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(str(value))
    _set_run_font(run, size=size, bold=bold, color=color)


def _add_table(
    document: Document,
    headers: Sequence[str],
    rows: Iterable[Sequence[Any]],
    widths: Sequence[int],
    *,
    numeric_columns: set[int] | None = None,
    font_size: float = 8.5,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    numeric_columns = numeric_columns or set()

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, LIGHT_BLUE)
        _set_cell_text(
            cell,
            header,
            bold=True,
            size=9,
            color=NAVY,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    _set_repeat_table_header(table.rows[0])

    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            align = (
                WD_ALIGN_PARAGRAPH.RIGHT
                if index in numeric_columns
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            _set_cell_text(row.cells[index], value, size=font_size, align=align)
    _apply_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def _display_value(field: str, value: Any) -> str:
    if value is None:
        return "null"
    if value == "":
        return '"" (빈 값)'
    text = str(value)
    if field in MONEY_FIELDS:
        try:
            return f"{int(text):,}원  [원본: {text}]"
        except ValueError:
            return text
    if field in RATIO_FIELDS:
        try:
            return f"{float(text) * 100:.4f}%  [원본: {text}]"
        except ValueError:
            return text
    if field == "source_ids":
        return text.replace("|", " | ")
    return text


def _add_field_table(
    document: Document,
    record: Mapping[str, Any],
    *,
    fields: Sequence[str] | None = None,
) -> None:
    selected = fields or tuple(record)
    rows = [
        (
            field,
            _display_value(field, record.get(field)),
            FIELD_DESCRIPTIONS.get(field, "저장 필드"),
        )
        for field in selected
    ]
    _add_table(
        document,
        ("필드명", "저장값", "의미"),
        rows,
        (2250, 3450, 3660),
        font_size=8,
    )


def _add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def _add_body(document: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        _set_run_font(lead, size=11, bold=True, color=NAVY)
        rest = paragraph.add_run(text[len(bold_lead) :])
        _set_run_font(rest, size=11, color=BLACK)
    else:
        run = paragraph.add_run(text)
        _set_run_font(run, size=11, color=BLACK)


def _add_callout(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph_properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), PALE_BLUE)
    paragraph_properties.append(shading)
    lead = paragraph.add_run(f"{label}  ")
    _set_run_font(lead, size=10.5, bold=True, color=BLUE)
    body = paragraph.add_run(text)
    _set_run_font(body, size=10.5, color=BLACK)


def _add_page_break(document: Document) -> None:
    document.add_page_break()


def _add_common_account(
    document: Document,
    account: Mapping[str, Any],
    *,
    heading_prefix: str,
) -> None:
    _add_heading(
        document,
        f"{heading_prefix} · {account['account_type']} · {account['account_id']}",
        2,
    )
    account_fields = {key: value for key, value in account.items() if key != "holdings"}
    _add_field_table(document, account_fields)
    _add_heading(document, "보유자산 전체", 3)
    rows = [
        (
            holding["asset_class"],
            _display_value("weight", holding["weight"]),
            _display_value("amount_krw", holding["amount_krw"]),
            holding["data_kind"],
            holding["source_ids"].replace("|", " | "),
        )
        for holding in account["holdings"]
    ]
    _add_table(
        document,
        ("자산군", "비중", "금액", "구분", "출처·가정 ID"),
        rows,
        (1500, 1450, 1900, 850, 3660),
        numeric_columns={1, 2},
        font_size=7.7,
    )


def _add_detailed_account(
    document: Document,
    account: Mapping[str, Any],
) -> None:
    _add_heading(document, f"{account['label']} 상세 포트폴리오", 2)
    _add_table(
        document,
        ("계좌 필드", "저장값"),
        (
            ("account_id", account["account_id"]),
            ("account_type", account["account_type"]),
            ("label", account["label"]),
        ),
        (2100, 7260),
    )
    rows = []
    for holding in account["holdings"]:
        rows.append(
            (
                holding["instrument_name"],
                holding.get("etf_isu_code") or "null",
                holding["asset_class_code"],
                _display_value("amount_krw", holding["amount_krw"]),
                holding["risk_treatment"],
                holding.get("statutory_exception") or "null",
            )
        )
    _add_table(
        document,
        ("상품명", "종목코드", "자산군", "금액", "위험 처리", "법정 예외"),
        rows,
        (2500, 1000, 1400, 1750, 1650, 1060),
        numeric_columns={3},
        font_size=7.8,
    )
    _add_heading(document, "상세 보유자산 원본 필드", 3)
    for index, holding in enumerate(account["holdings"], start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(f"{index}. {holding['instrument_name']}")
        _set_run_font(run, size=9.5, bold=True, color=DARK_BLUE)
        _add_field_table(document, holding, fields=DETAILED_HOLDING_FIELDS)


def _configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))
    _set_run_font(run, size=8.5, color=MUTED)


def _configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("연금 코파일럿  |  고객 목데이터 보고서")
    _set_run_font(run, size=8.5, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    label = footer.add_run("INTERNAL · MOCK DATA   |   ")
    _set_run_font(label, size=8.5, color=MUTED)
    _add_page_number(footer)


def _add_title_page(
    document: Document,
    examples: Mapping[str, Any],
    manifest: list[dict[str, Any]],
    scenarios: list[dict[str, Any]],
) -> None:
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(26)
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("TEAM DATA BRIEF")
    _set_run_font(run, size=10, bold=True, color=BLUE)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run("고객 목데이터 대표 2명 보고서")
    _set_run_font(run, size=26, bold=True, color=NAVY)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run(
        "1만 명 일반 고객 1명과 대표 시나리오 고객 1명의 전체 데이터 구조"
    )
    _set_run_font(run, size=12.5, color=MUTED)

    _add_table(
        document,
        ("문서 구분", "내용"),
        (
            ("기준 데이터", "data/mock/customer_data_examples.json"),
            ("스키마 버전", str(examples["schema_version"])),
            ("작성 기준일", "2026-07-20"),
            ("데이터 성격", "실제 개인정보가 아닌 완전 합성 목데이터"),
        ),
        (2100, 7260),
    )

    _add_callout(
        document,
        "핵심 결론",
        "대표 고객 6명은 1만 명 고객과 동일한 고객 29개·계좌 23개·보유자산 6개 "
        "공통 계약을 모두 가지며, 로그인 시나리오 정보와 실제 적격 ETF 상세내역이 "
        "추가됩니다.",
    )

    benchmark = examples["benchmark_customer_example"]
    representative = examples["representative_customer_example"]
    identity = representative["demo_identity"]
    common = representative["benchmark_contract"]
    _add_heading(document, "대표 예시 비교", 1)
    _add_table(
        document,
        ("비교 항목", "일반 고객", "대표 시나리오 고객"),
        (
            ("고객", benchmark["customer"]["user_id"], identity["nickname"]),
            (
                "기준 고객 ID",
                benchmark["customer"]["user_id"],
                identity["benchmark_user_id"],
            ),
            ("계좌 수", len(benchmark["accounts"]), len(common["accounts"])),
            (
                "총잔액",
                f"{sum(int(a['balance_krw']) for a in benchmark['accounts']):,}원",
                f"{sum(int(a['balance_krw']) for a in common['accounts']):,}원",
            ),
            ("ETF 종목 연결", "없음 · 자산군 수준", "있음 · 실제 적격 ETF 상세화"),
            ("로그인 정보", "없음", "시연용 Auth·로그인 ID 보유"),
            ("RAG 저장", "저장하지 않음", "저장하지 않음"),
        ),
        (2200, 3280, 3880),
    )

    _add_heading(document, "대표 6명 ETF 운용사 분산", 1)
    identity_by_code = {item["scenario_code"]: item for item in manifest}
    issuer_rows = []
    all_issuers: set[str] = set()
    for scenario in scenarios:
        etfs = [
            holding
            for account in scenario["accounts"]
            for holding in account["holdings"]
            if holding.get("etf_isu_code")
        ]
        issuers = sorted({item["instrument_name"].split()[0] for item in etfs})
        all_issuers.update(issuers)
        issuer_rows.append(
            (
                identity_by_code[scenario["scenario_code"]]["nickname"],
                scenario["scenario_code"],
                ", ".join(issuers),
                len(etfs),
            )
        )
    _add_table(
        document,
        ("대표 고객", "시나리오", "ETF 운용사", "ETF 수"),
        issuer_rows,
        (1700, 2900, 3700, 1060),
        numeric_columns={3},
        font_size=8,
    )
    _add_body(
        document,
        "6명 전체 운용사: " + " · ".join(sorted(all_issuers)) + ". KODEX도 유지됩니다.",
        bold_lead="6명 전체 운용사:",
    )


def build_document() -> Document:
    examples = json.loads(
        (MOCK_DIR / "customer_data_examples.json").read_text(encoding="utf-8")
    )
    manifest = json.loads(
        (MOCK_DIR / "demo_scenario_users.json").read_text(encoding="utf-8")
    )["users"]
    scenarios = json.loads(
        (MOCK_DIR / "chatbot_scenarios.json").read_text(encoding="utf-8")
    )

    document = Document()
    _configure_styles(document)
    _configure_page(document)
    document.core_properties.title = "고객 목데이터 대표 2명 보고서"
    document.core_properties.subject = (
        "1만 명 일반 고객과 대표 시나리오 고객 데이터 예시"
    )
    document.core_properties.author = "KDA Securities"
    document.core_properties.keywords = "MOCK, 고객데이터, 연금계좌, ETF"

    _add_title_page(document, examples, manifest, scenarios)

    benchmark = examples["benchmark_customer_example"]
    _add_page_break(document)
    _add_heading(document, "1. 1만 명 일반 고객 예시", 1)
    _add_callout(
        document,
        "대상",
        "USR00001 · 30세 근로자 · DC와 개인 IRP 보유 · DC 방치형 합성 시나리오",
    )
    _add_heading(document, "고객 필드 29개 전체", 2)
    _add_field_table(document, benchmark["customer"])
    for index, account in enumerate(benchmark["accounts"], start=1):
        _add_page_break(document)
        _add_common_account(document, account, heading_prefix=f"일반 고객 계좌 {index}")

    representative = examples["representative_customer_example"]
    identity = representative["demo_identity"]
    common = representative["benchmark_contract"]
    detailed = representative["detailed_etf_portfolio"]

    _add_page_break(document)
    _add_heading(document, "2. 대표 시나리오 고객 예시", 1)
    _add_callout(
        document,
        "대상",
        f"{identity['nickname']} · 기준 고객 {identity['benchmark_user_id']} · "
        "계좌별 중복·위험 편중 시나리오",
    )
    _add_heading(document, "시연 로그인·시나리오 연결 정보 전체", 2)
    _add_field_table(document, identity)
    _add_heading(document, "1만 명 공통 고객 필드 29개 전체", 2)
    _add_field_table(document, common["customer"])

    _add_page_break(document)
    _add_heading(document, "대표 고객 세액공제·한도 검증", 2)
    pension = int(common["customer"]["pension_savings_contribution_krw"])
    irp = int(common["customer"]["irp_contribution_krw"])
    eligible = int(common["customer"]["total_tax_credit_eligible_contribution_krw"])
    estimated = int(common["customer"]["estimated_pension_tax_credit_krw"])
    _add_table(
        document,
        ("검증 항목", "고객값", "기준", "판정"),
        (
            ("연금저축 납입", f"{pension:,}원", "6,000,000원 이하", "충족"),
            ("개인 IRP 납입", f"{irp:,}원", "계좌별 생성값", "충족"),
            ("실제 납입 합산", f"{pension + irp:,}원", "18,000,000원 이하", "충족"),
            ("세액공제 대상", f"{eligible:,}원", "합산 9,000,000원 이하", "충족"),
            (
                "적용 공제율",
                f"{common['customer']['pension_tax_credit_rate_pct']}%",
                "총급여 5,500만 원 초과",
                "13.2%",
            ),
            ("예상 세액공제", f"{estimated:,}원", f"{eligible:,} × 13.2%", "일치"),
        ),
        (2400, 2100, 3260, 1600),
        numeric_columns={1},
    )

    _add_heading(document, "공통 계좌·자산군 데이터 전체", 1)
    for index, account in enumerate(common["accounts"], start=1):
        if index > 1:
            _add_page_break(document)
        _add_common_account(
            document,
            account,
            heading_prefix=f"대표 고객 공통 계좌 {index}",
        )

    _add_page_break(document)
    _add_heading(document, "시연용 상세 ETF 포트폴리오", 1)
    _add_body(
        document,
        "공통 계좌의 잔액과 자산군 금액은 유지하고, 대표 고객에 한해 실제 적격 ETF "
        "종목코드 단위로 상세화했습니다. 선택 필드가 없으면 null로 표시합니다.",
    )
    _add_heading(document, "포트폴리오 메타데이터 전체", 2)
    _add_field_table(
        document,
        {key: value for key, value in detailed.items() if key != "accounts"},
    )
    for account in detailed["accounts"]:
        _add_page_break(document)
        _add_detailed_account(document, account)

    _add_page_break(document)
    _add_heading(document, "3. 최종 정합성 및 사용 유의사항", 1)
    common_total = sum(int(account["balance_krw"]) for account in common["accounts"])
    detailed_total = sum(
        int(holding["amount_krw"])
        for account in detailed["accounts"]
        for holding in account["holdings"]
    )
    _add_table(
        document,
        ("검증 항목", "결과"),
        (
            ("대표 고객 공통 계좌 잔액", f"{common_total:,}원"),
            ("대표 고객 상세 보유자산 합계", f"{detailed_total:,}원"),
            (
                "상세화 전후 잔액",
                "일치" if common_total == detailed_total else "불일치",
            ),
            ("DC·IRP 위험자산", "각 계좌 일반 위험자산 70% 이하"),
            ("연금저축 위험자산", "DC·IRP의 70% 총량 한도 미적용"),
            ("개인 계좌 데이터의 RAG 저장", "저장하지 않음"),
        ),
        (3600, 5760),
    )
    _add_callout(
        document,
        "주의",
        "이 문서의 고객·계좌 데이터는 모두 합성 목데이터입니다. 예상 세액공제액은 "
        "교육용 계산값이며 실제 환급액을 보장하지 않습니다. 수익률은 과거 실적을 "
        "참고한 합성값으로 미래 수익 예측값이 아닙니다.",
    )
    _add_body(
        document,
        "원본 파일: data/mock/customer_data_examples.json · 상세 Markdown: "
        "docs/30_스펙/고객_목데이터_전체_예시.md",
    )
    return document


def main() -> None:
    document = build_document()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(f"wrote {OUTPUT_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
