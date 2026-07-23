"""Build the gitignored six-customer demo report including Auth credentials."""

from __future__ import annotations

import json
import zipfile
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt

from scripts.build_customer_data_report_docx import (
    BLUE,
    MUTED,
    NAVY,
    ROOT,
    _add_body,
    _add_callout,
    _add_heading,
    _add_page_break,
    _add_table,
    _configure_page,
    _configure_styles,
    _set_run_font,
)
from scripts.build_demo_customer_portfolios import _complete_customer

MOCK_DIR = ROOT / "data" / "mock"
CREDENTIALS_PATH = ROOT / "secrets" / "demo_scenario_auth.json"
OUTPUT_PATH = (
    ROOT / "secrets" / "대표_시나리오_고객_6명_아이디_비밀번호_포함_보고서.docx"
)
ASSET_LABELS = {
    "cash": "현금성",
    "deposit": "원리금보장",
    "bond": "채권",
    "domestic_equity": "국내주식",
    "global_equity": "글로벌주식",
}
SOURCE_LOCATION_ROWS = (
    (
        "로그인 ID·비밀번호",
        "secrets/demo_scenario_auth.json",
        "users[].scenario_code / login_id / auth_email / password",
        "화면은 짧은 login_id 사용 · Auth는 내부 auth_email 사용",
    ),
    (
        "대표 고객 식별·시연 후보",
        "data/mock/demo_scenario_users.json",
        "users[].scenario_code / auth_user_id / benchmark_user_id",
        "auth.users.app_metadata · public.demo_user_financial_context",
    ),
    (
        "고객·소득·세액공제",
        "data/mock/users.csv",
        "user_id = benchmark_user_id",
        "benchmark.benchmark_mock_users · public.demo_user_financial_context",
    ),
    (
        "계좌 요약·연간 납입",
        "data/mock/accounts.csv",
        "user_id = benchmark_user_id",
        "benchmark.benchmark_mock_accounts · public.mock_accounts",
    ),
    (
        "1만 명 기준 보유자산",
        "data/mock/holdings.csv",
        "account_id = accounts.csv의 account_id",
        "benchmark.benchmark_mock_holdings",
    ),
    (
        "투자성향·배점·이유·후기",
        "data/mock/demo_investor_profiles.json",
        "profiles[].scenario_code",
        "public.demo_investor_profiles · demo_investor_profile_answers",
    ),
    (
        "대표 고객 상세 ETF 비중",
        "data/mock/chatbot_scenarios.json",
        "[].scenario_code / accounts[].holdings[]",
        "public.mock_accounts · public.mock_holdings",
    ),
    (
        "과거 수익률·추천(좋아요)",
        "data/mock/demo_public_portfolio_metrics.json",
        "profiles[].scenario_code",
        "public.demo_public_portfolio_metrics · 인증 GET /chat/heroes",
    ),
    (
        "ETF 상품·과거수익 근거",
        "ETF 정규화 산출물 및 검증 스크립트",
        "etf_isu_code / isu_code",
        "public.etf_universe_products · public.etf_return_histories",
    ),
    (
        "설계·필드 계약",
        "docs/30_스펙/대표_시나리오_투자성향_포트폴리오_계약.md",
        "성향·포트폴리오 SSOT와 표시 계약",
        "DB_HANDOFF.md에서 원격 적용 이력 확인",
    ),
)


def _load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def _validate_private_credentials(
    users: list[dict[str, Any]], credentials: list[dict[str, str]]
) -> None:
    expected = {
        (
            str(user["auth_user_id"]),
            str(user["scenario_code"]),
            str(user["login_id"]),
            str(user["auth_email"]),
        )
        for user in users
    }
    actual = {
        (
            str(item["auth_user_id"]),
            str(item["scenario_code"]),
            str(item["login_id"]),
            str(item["auth_email"]),
        )
        for item in credentials
    }
    passwords = [str(item["password"]) for item in credentials]
    if expected != actual or len(credentials) != 6:
        raise ValueError("private credentials do not match the six-customer manifest")
    if len(set(passwords)) != 6 or any(len(password) < 6 for password in passwords):
        raise ValueError("private credentials must contain six valid unique passwords")


def _money(value: Any) -> str:
    if value in {None, ""}:
        return "-"
    return f"{int(value):,}원"


def _value(value: Any) -> str:
    if value in {None, ""}:
        return "-"
    if isinstance(value, bool):
        return "예" if value else "아니요"
    return str(value)


def _percent(amount: int, total: int) -> str:
    if total == 0:
        return "0.00%"
    value = (Decimal(amount) / Decimal(total) * Decimal("100")).quantize(
        Decimal("0.01"), rounding=ROUND_HALF_UP
    )
    return f"{value}%"


def _configure_private_report(document: Document) -> None:
    _configure_styles(document)
    _configure_page(document)
    header = document.sections[0].header.paragraphs[0]
    header.text = ""
    run = header.add_run("연금 코파일럿  |  대표 고객 6명 비공개 시연 자료")
    _set_run_font(run, size=8.5, bold=True, color=MUTED)


def _add_title_page(
    document: Document,
    users: list[dict[str, Any]],
    credentials_by_code: dict[str, dict[str, str]],
    profiles_by_code: dict[str, dict[str, Any]],
) -> None:
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(22)
    kicker.paragraph_format.space_after = Pt(4)
    _set_run_font(
        kicker.add_run("INTERNAL DEMO CREDENTIALS"),
        size=10,
        bold=True,
        color=BLUE,
    )

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    _set_run_font(
        title.add_run("대표 시나리오 고객 6명 데이터 리포트"),
        size=24,
        bold=True,
        color=NAVY,
    )

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    _set_run_font(
        subtitle.add_run("로그인 아이디·임시 고정 비밀번호 포함"),
        size=12.5,
        color=MUTED,
    )

    _add_callout(
        document,
        "외부 공유 금지",
        "이 문서는 발표·학습용 합성 고객의 로그인 자격증명을 평문으로 포함합니다. "
        "Git·메신저 공개 채널·화면 녹화본에 첨부하지 마세요.",
    )
    _add_callout(
        document,
        "비밀번호 안내",
        "Supabase Auth 최소 길이가 6자이므로 요청한 4자 값 대신 같은 문자를 "
        "6번 반복한 값을 실제 로그인 비밀번호로 적용했습니다. 기존 강한 비밀번호 "
        "생성 로직은 유지되며, 현재 비공개 credentials 파일이 존재하는 동안 "
        "재생성되지 않습니다.",
    )

    credential_rows = []
    for index, user in enumerate(users, start=1):
        profile = profiles_by_code[str(user["scenario_code"])]
        credential = credentials_by_code[str(user["scenario_code"])]
        candidate = "후보" if user["is_demo_login_candidate"] else "후보 제외"
        credential_rows.append(
            (
                index,
                user["nickname"],
                profile["investor_profile_label"],
                candidate,
                credential["login_id"],
                credential["password"],
            )
        )
    _add_heading(document, "로그인 자격증명 요약", 1)
    _add_table(
        document,
        ("번호", "대표 고객", "투자성향", "시연", "로그인 ID", "비밀번호"),
        credential_rows,
        (600, 1400, 1200, 1100, 3260, 1800),
        numeric_columns={0},
        font_size=8.2,
    )
    _add_body(
        document,
        "6번 윤정희 고객은 데이터와 Auth 계정을 유지하지만 당일 시연 로그인 "
        "후보에서는 제외합니다.",
        bold_lead="6번 윤정희 고객은",
    )


def _add_source_location_guide(
    document: Document,
    users: list[dict[str, Any]],
    metrics_by_code: dict[str, dict[str, Any]],
    metric_payload: dict[str, Any],
) -> None:
    _add_page_break(document)
    _add_heading(document, "정보별 저장 위치 안내", 1)
    _add_callout(
        document,
        "찾는 순서",
        "먼저 scenario_code로 대표 고객을 찾고, benchmark_user_id로 1만 명 기준 "
        "고객·계좌 데이터를 연결합니다. 로그인 계정은 auth_user_id로 Supabase "
        "Auth와 금융 컨텍스트를 조회합니다.",
    )
    _add_table(
        document,
        ("정보", "로컬 기준 위치", "조회 키·필드", "Supabase·런타임 위치"),
        SOURCE_LOCATION_ROWS,
        (1500, 2920, 2220, 2720),
        font_size=7.2,
    )
    _add_heading(document, "대표 고객 공개 비교 지표", 2)
    _add_table(
        document,
        ("번호", "고객", "과거 12개월 수익률", "산정 기간", "좋아요", "데이터 성격"),
        [
            (
                index,
                user["nickname"],
                f"{metrics_by_code[str(user['scenario_code'])]['portfolio_trailing_12m_return_pct']}%",
                (
                    f"{metrics_by_code[str(user['scenario_code'])]['return_period_start']}~"
                    f"{metrics_by_code[str(user['scenario_code'])]['return_period_end']}"
                ),
                metrics_by_code[str(user["scenario_code"])]["like_count"],
                "수익률·좋아요 모두 MOCK",
            )
            for index, user in enumerate(users, start=1)
        ],
        (600, 1400, 1450, 2200, 900, 2810),
        numeric_columns={0, 2, 4},
        font_size=7.6,
    )
    _add_body(
        document,
        f"수익률 계산: {metric_payload['return_metric']['calculation_basis']}. "
        "공식 커뮤니티 랭킹 산식이 아니라 화면 표시용 참고값이며 미래 예측이 "
        "아닙니다. 좋아요 수는 수익률과 무관하게 배정했습니다.",
        bold_lead="수익률 계산:",
    )
    _add_body(
        document,
        "비밀번호 평문은 Git 제외 secrets 파일과 이 비공개 보고서에만 있습니다. "
        "투자성향·수기 투자 이유·후기의 현재 SSOT는 "
        "data/mock/demo_investor_profiles.json이며, ETF 상세 보유의 SSOT는 "
        "data/mock/chatbot_scenarios.json입니다.",
        bold_lead="비밀번호 평문은",
    )


def _add_customer_overview(
    document: Document,
    *,
    index: int,
    user: dict[str, Any],
    credential: dict[str, str],
    profile: dict[str, Any],
    customer_record: dict[str, Any],
    scenario: dict[str, Any],
    metric: dict[str, Any],
    metric_payload: dict[str, Any],
) -> None:
    customer = customer_record["customer"]
    accounts = customer_record["accounts"]
    total_balance = sum(int(account["balance_krw"]) for account in accounts)
    candidate = "시연 로그인 후보" if user["is_demo_login_candidate"] else "후보 제외"

    _add_page_break(document)
    _add_heading(
        document,
        f"{index}. {user['nickname']} · {profile['investor_profile_label']}",
        1,
    )
    _add_callout(document, "대표 상황", profile["scenario_description"])
    _add_table(
        document,
        ("고객별 조회 키", "값", "주요 확인 위치", "연결 기준"),
        (
            (
                "scenario_code",
                user["scenario_code"],
                "demo_scenario_users.json · demo_investor_profiles.json · "
                "chatbot_scenarios.json",
                "세 파일의 scenario_code",
            ),
            (
                "benchmark_user_id",
                user["benchmark_user_id"],
                "users.csv · accounts.csv · holdings.csv",
                "users.csv user_id",
            ),
            (
                "auth_user_id",
                user["auth_user_id"],
                "secrets/demo_scenario_auth.json · Supabase auth.users",
                "Auth id / demo_user_financial_context.user_id",
            ),
        ),
        (1500, 2180, 3440, 2240),
        font_size=7.4,
    )

    _add_heading(document, "로그인·식별 정보", 2)
    _add_table(
        document,
        ("항목", "값", "항목", "값"),
        (
            ("로그인 ID", credential["login_id"], "비밀번호", credential["password"]),
            ("시연 구분", candidate, "대표 나이", f"{user['representative_age']}세"),
            ("시나리오", user["scenario_code"], "기준 고객", user["benchmark_user_id"]),
            ("Auth UUID", user["auth_user_id"], "데이터 구분", "MOCK"),
        ),
        (1500, 3180, 1500, 3180),
        font_size=8.2,
    )

    _add_heading(document, "고객·소득·세액공제 데이터", 2)
    _add_table(
        document,
        ("항목", "값", "항목", "값"),
        (
            (
                "고용 형태",
                customer["employment_type"],
                "총 계좌잔액",
                _money(total_balance),
            ),
            (
                "총급여액",
                _money(customer["gross_salary_krw"]),
                "종합소득금액",
                _money(customer["comprehensive_income_krw"]),
            ),
            (
                "귀속연도",
                customer["tax_year"],
                "세액공제율",
                f"{customer['pension_tax_credit_rate_pct']}%",
            ),
            (
                "연금저축 납입",
                _money(customer["pension_savings_contribution_krw"]),
                "개인 IRP 납입",
                _money(customer["irp_contribution_krw"]),
            ),
            (
                "공제대상 합계",
                _money(customer["total_tax_credit_eligible_contribution_krw"]),
                "예상 세액공제",
                _money(customer["estimated_pension_tax_credit_krw"]),
            ),
            (
                "연금 시작 나이",
                f"{customer['planned_pension_start_age']}세",
                "수령 선호",
                customer["payout_preference"],
            ),
        ),
        (1500, 3180, 1500, 3180),
        font_size=8.1,
    )

    _add_heading(document, "투자성향·공개 포트폴리오 문구", 2)
    _add_table(
        document,
        ("항목", "저장값"),
        (
            (
                "투자성향",
                f"{profile['investor_profile_label']} "
                f"({profile['total_score']}점 · {profile['score_band']})",
            ),
            ("투자 이유", profile["investment_reason"]),
            ("투자의견·후기", profile["portfolio_opinion_review"]),
            (
                "대표 ETF 종목코드",
                ", ".join(profile["representative_etf_isu_codes"]),
            ),
            ("대표 ETF 테마", profile["representative_etf_theme"]),
            ("대표 ETF 테마 후기", profile["representative_etf_theme_review"]),
            ("포트폴리오 정합성", profile["portfolio_consistency_note"]),
        ),
        (1900, 7460),
        font_size=8.4,
    )

    _add_heading(document, "과거 수익률·추천(좋아요) 지표", 2)
    _add_table(
        document,
        ("항목", "값", "항목", "값"),
        (
            (
                metric_payload["return_metric"]["label"],
                f"{metric['portfolio_trailing_12m_return_pct']}%",
                "산정 기간",
                f"{metric['return_period_start']}~{metric['return_period_end']}",
            ),
            (
                "계산 기준",
                metric_payload["return_metric"]["calculation_basis"],
                "미래 예측",
                "아니요",
            ),
            (
                metric_payload["like_metric"]["label"],
                f"{int(metric['like_count']):,}개",
                "좋아요 기준일",
                metric_payload["like_metric"]["as_of_date"],
            ),
            (
                "데이터 성격",
                "MOCK",
                "공식 랭킹 사용",
                "아니요",
            ),
        ),
        (1500, 3180, 1500, 3180),
        font_size=7.7,
    )

    _add_heading(document, "투자성향 배점 11개", 2)
    _add_table(
        document,
        ("문항", "선택 응답", "점수", "판단 근거"),
        [
            (
                answer["question_code"],
                answer["selected_option"],
                answer["score"],
                answer["basis"],
            )
            for answer in profile["answers"]
        ],
        (2000, 3000, 700, 3660),
        numeric_columns={2},
        font_size=7.5,
    )
    _add_body(
        document,
        "비채점 응답: 파생상품 경험 "
        f"{profile['non_scored_answers']['derivative_experience']} · "
        "취약 금융소비자 "
        f"{_value(profile['non_scored_answers']['vulnerable_financial_consumer'])} · "
        "유효기간 동의 "
        f"{_value(profile['non_scored_answers']['assessment_validity_consent'])}",
    )

    _add_heading(document, "계좌 요약", 2)
    _add_table(
        document,
        ("계좌", "잔액", "연간 납입", "위험자산", "현금성", "수령 예상액"),
        [
            (
                account["account_type"],
                _money(account["balance_krw"]),
                _money(account["annual_contribution_krw"]),
                f"{float(account['risky_asset_ratio']) * 100:.2f}%",
                f"{float(account['cash_ratio']) * 100:.2f}%",
                _money(account["planned_annual_pension_receipt_krw"]),
            )
            for account in accounts
        ],
        (1250, 1700, 1600, 1500, 1400, 1910),
        numeric_columns={1, 2, 3, 4, 5},
        font_size=7.8,
    )

    detailed_accounts = {item["account_type"]: item for item in scenario["accounts"]}
    account_balance_by_type = {
        account["account_type"]
        .lower()
        .replace("pension_savings_fund", "pension_savings"): int(
            account["balance_krw"]
        )
        for account in accounts
    }
    holding_rows = []
    for account_type, account in detailed_accounts.items():
        total = account_balance_by_type[account_type]
        for holding in account["holdings"]:
            amount = int(holding["amount_krw"])
            holding_rows.append(
                (
                    account["label"],
                    holding["instrument_name"],
                    holding.get("etf_isu_code") or "-",
                    ASSET_LABELS.get(
                        holding["asset_class_code"], holding["asset_class_code"]
                    ),
                    _money(amount),
                    _percent(amount, total),
                )
            )
    _add_heading(document, "성향 기반 계좌별 상세 ETF·자산", 2)
    _add_table(
        document,
        ("계좌", "상품·자산", "종목코드", "자산군", "금액", "계좌 비중"),
        holding_rows,
        (1350, 2600, 1000, 1300, 1650, 1460),
        numeric_columns={4, 5},
        font_size=7.5,
    )


def build_document() -> Document:
    users = _load_json(MOCK_DIR / "demo_scenario_users.json")["users"]
    credentials_payload = _load_json(CREDENTIALS_PATH)
    credentials = credentials_payload["users"]
    _validate_private_credentials(users, credentials)
    credentials_by_code = {item["scenario_code"]: item for item in credentials}

    profile_payload = _load_json(MOCK_DIR / "demo_investor_profiles.json")
    profiles_by_code = {
        item["scenario_code"]: item for item in profile_payload["profiles"]
    }
    scenarios = _load_json(MOCK_DIR / "chatbot_scenarios.json")
    scenarios_by_code = {item["scenario_code"]: item for item in scenarios}
    metric_payload = _load_json(MOCK_DIR / "demo_public_portfolio_metrics.json")
    metrics_by_code = {
        item["scenario_code"]: item for item in metric_payload["profiles"]
    }
    if set(credentials_by_code) != set(profiles_by_code) or set(
        profiles_by_code
    ) != set(scenarios_by_code) or set(scenarios_by_code) != set(metrics_by_code):
        raise ValueError(
            "credentials, profiles, scenarios, and metrics must contain the same users"
        )

    document = Document()
    _configure_private_report(document)
    document.core_properties.title = "대표 시나리오 고객 6명 데이터 리포트"
    document.core_properties.subject = "시연 로그인 자격증명 포함 비공개 합성 고객 자료"
    document.core_properties.author = "연금 코파일럿 팀"
    _add_title_page(document, users, credentials_by_code, profiles_by_code)
    _add_source_location_guide(document, users, metrics_by_code, metric_payload)

    for index, user in enumerate(users, start=1):
        code = str(user["scenario_code"])
        _add_customer_overview(
            document,
            index=index,
            user=user,
            credential=credentials_by_code[code],
            profile=profiles_by_code[code],
            customer_record=_complete_customer(str(user["benchmark_user_id"])),
            scenario=scenarios_by_code[code],
            metric=metrics_by_code[code],
            metric_payload=metric_payload,
        )

    _add_page_break(document)
    _add_heading(document, "사용 유의사항", 1)
    _add_callout(
        document,
        "중요",
        "모든 고객·소득·계좌·투자성향·후기는 합성 목데이터입니다. "
        "실제 고객, 미래 수익률 예측 또는 투자성과 보장으로 사용하지 않습니다.",
    )
    _add_body(
        document,
        "비밀번호를 다시 강한 임시값으로 전환할 때는 비공개 credentials 파일을 "
        "안전하게 교체한 뒤 기존 provision_demo_auth_users.py의 "
        "--rotate-existing 경로로 6명 로그인을 다시 검증합니다.",
    )
    return document


def _verify_saved_report() -> None:
    with zipfile.ZipFile(OUTPUT_PATH) as archive:
        if archive.testzip() is not None:
            raise ValueError("private report DOCX ZIP is corrupt")

    document = Document(OUTPUT_PATH)
    if len(document.tables) != 51:
        table_count = len(document.tables)
        raise ValueError(f"unexpected private report table count: {table_count}")

    text_parts = [paragraph.text for paragraph in document.paragraphs]
    text_parts.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    report_text = "\n".join(text_parts)

    credentials = _load_json(CREDENTIALS_PATH)["users"]
    for credential in credentials:
        if credential["login_id"] not in report_text:
            raise ValueError("a demo login ID is missing from the private report")
        if credential["password"] not in report_text:
            raise ValueError("a demo password is missing from the private report")

    metrics = _load_json(MOCK_DIR / "demo_public_portfolio_metrics.json")
    required_labels = (
        "대표 고객 공개 비교 지표",
        "과거 수익률·추천(좋아요) 지표",
        "공식 랭킹 사용",
        "data/mock/demo_public_portfolio_metrics.json",
    )
    if any(label not in report_text for label in required_labels):
        raise ValueError("public portfolio metric disclosure is incomplete")
    for metric in metrics["profiles"]:
        if f"{metric['portfolio_trailing_12m_return_pct']}%" not in report_text:
            raise ValueError("a past return is missing from the private report")
        if f"{int(metric['like_count']):,}개" not in report_text:
            raise ValueError("a like count is missing from the private report")
    if "{{" in report_text or "}}" in report_text:
        raise ValueError("template placeholder remains in the private report")

    section = document.sections[0]
    expected_dimensions = {
        "page_width": Inches(8.5),
        "page_height": Inches(11),
        "top_margin": Inches(1),
        "right_margin": Inches(1),
        "bottom_margin": Inches(1),
        "left_margin": Inches(1),
    }
    for attribute, expected in expected_dimensions.items():
        if getattr(section, attribute) != expected:
            raise ValueError(f"unexpected private report page setting: {attribute}")


def main() -> None:
    document = build_document()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    _verify_saved_report()
    print(f"wrote private report: {OUTPUT_PATH.name}")


if __name__ == "__main__":
    main()
