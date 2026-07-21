"""Refresh audit metadata in the six-customer demo report."""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from scripts.validate_demo_scenario_etfs import validate

ROOT = Path(__file__).resolve().parents[1]
REPORT_PATH = (
    ROOT / "docs" / "30_스펙" / "대표_시나리오_고객_6명_시연용_고객정보.docx"
)


def _set_cell(cell, text: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(8)


def _set_row(row, values: tuple[str, ...]) -> None:
    if len(row.cells) != len(values):
        raise ValueError(f"row has {len(row.cells)} cells, expected {len(values)}")
    for cell, value in zip(row.cells, values, strict=True):
        _set_cell(cell, value)


def _upsert_row(table, key: str, values: tuple[str, ...]) -> None:
    for row in table.rows:
        if row.cells[0].text == key:
            _set_row(row, values)
            return
    _set_row(table.add_row(), values)


def _remove_manual_page_breaks(document) -> None:
    for paragraph in document.paragraphs:
        for page_break in paragraph._element.xpath('.//w:br[@w:type="page"]'):
            page_break.getparent().remove(page_break)


def refresh() -> None:
    result = validate()
    if result["validation_error_count"]:
        raise ValueError(f"ETF validation failed: {result['errors']}")
    if result["scenario_count"] != 6 or result["account_count"] != 13:
        raise ValueError("unexpected demo customer or account count")

    document = Document(REPORT_PATH)
    if len(document.tables) != 64:
        raise ValueError(f"unexpected table count: {len(document.tables)}")

    metadata = document.tables[0]
    _set_row(
        metadata.rows[0],
        (
            "작성 기준일",
            "2026-07-21",
            "데이터 구분",
            "발표·학습용 가상 목데이터",
        ),
    )
    _set_row(
        metadata.rows[1],
        ("대표 고객", "6명", "전체 상세 보유행", "86건"),
    )
    _upsert_row(
        metadata,
        "ETF 보유행·고유 종목",
        (
            "ETF 보유행·고유 종목",
            f"{result['etf_holding_count']}건 · {result['unique_etf_count']}종",
            "운용사",
            "ACE·HANARO·KODEX·RISE·SOL·TIGER",
        ),
    )
    _upsert_row(
        metadata,
        "ETF 재검증",
        (
            "ETF 재검증",
            "계좌별 적격 65/65",
            "시장 근거",
            "253개 관측 충족 65/65 · 위험한도 위반 0",
        ),
    )

    login_table = document.tables[2]
    _set_cell(login_table.rows[6].cells[0], "6 (후보 제외)")
    _set_cell(login_table.rows[6].cells[3], "시연 로그인 후보 제외")
    _set_cell(
        login_table.rows[6].cells[4],
        "연금 수령 전환 · 데이터/Auth 계정 유지",
    )

    operation_table = document.tables[3]
    _set_cell(
        operation_table.rows[1].cells[1],
        "시연 로그인 후보 5명 중 당일 사용할 1명을 시연자가 임의로 선택",
    )

    payout_login_table = document.tables[54]
    _set_cell(payout_login_table.rows[0].cells[3], "시연 로그인 후보 제외")

    checklist = document.tables[62]
    _set_cell(
        checklist.rows[1].cells[1],
        "시연 로그인 후보 5명 중 사용할 1명을 정하고 로그인 ID 확인",
    )
    _upsert_row(
        checklist,
        "6번 고객",
        (
            "6번 고객",
            "데이터·Auth 계정과 타 고객 포트폴리오는 유지하되 로그인 후보에서 제외",
        ),
    )
    _upsert_row(
        checklist,
        "ETF 검증",
        (
            "ETF 검증",
            "65개 ETF 보유행의 계좌 적격성·253개 과거 관측·DC/IRP 70% 한도 확인",
        ),
    )

    _set_cell(
        document.tables[63].rows[0].cells[0],
        "발표 시 한 문장\n"
        "시연 로그인 후보 5명 중 한 명으로 로그인하면 본인은 제외되고, "
        "나머지 대표 고객 5명의 상세 포트폴리오를 비교한 뒤 내 계좌 규칙에 "
        "맞는 비중 적용 시뮬레이션을 확인할 수 있습니다.",
    )

    _remove_manual_page_breaks(document)
    temporary_path = REPORT_PATH.with_suffix(".tmp.docx")
    document.save(temporary_path)
    os.replace(temporary_path, REPORT_PATH)


def main() -> None:
    refresh()
    print(REPORT_PATH)


if __name__ == "__main__":
    main()
