"""Build a concise team-facing Word summary for two mock customer examples."""

from __future__ import annotations

import json
from collections.abc import Mapping
from typing import Any

from docx import Document
from docx.shared import Pt

from scripts.build_customer_data_report_docx import (
    BLUE,
    MOCK_DIR,
    MUTED,
    NAVY,
    ROOT,
    _add_body,
    _add_callout,
    _add_heading,
    _add_page_break,
    _add_table,
    _configure_page,
    _configure_styles,
    _set_run_font,
)

OUTPUT_PATH = ROOT / "docs" / "30_스펙" / "고객_목데이터_대표_2명_요약본.docx"


def _money(value: Any) -> str:
    return f"{int(value):,}원"


def _percent(value: Any, *, ratio: bool = False) -> str:
    number = float(value)
    if ratio:
        number *= 100
    return f"{number:.2f}%"


def _account_total(accounts: list[Mapping[str, Any]]) -> int:
    return sum(int(account["balance_krw"]) for account in accounts)


def _issuer_name(instrument_name: str) -> str | None:
    issuer = instrument_name.split()[0]
    return (
        issuer if issuer in {"KODEX", "TIGER", "ACE", "RISE", "SOL", "HANARO"} else None
    )


def _configure_summary_page(document: Document) -> None:
    _configure_page(document)
    header = document.sections[0].header.paragraphs[0]
    header.clear()
    run = header.add_run("연금 코파일럿  |  고객 목데이터 요약")
    _set_run_font(run, size=8.5, bold=True, color=MUTED)


def _add_title(document: Document) -> None:
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(3)
    _set_run_font(
        kicker.add_run("TEAM EXECUTIVE SUMMARY"), size=9.5, bold=True, color=BLUE
    )

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    _set_run_font(
        title.add_run("고객 목데이터 대표 2명 요약"), size=24, bold=True, color=NAVY
    )

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)
    _set_run_font(
        subtitle.add_run("1만 명 일반 고객과 대표 시나리오 고객의 저장 데이터 비교"),
        size=12,
        color=MUTED,
    )


def _add_overview(
    document: Document,
    examples: Mapping[str, Any],
    scenarios: list[Mapping[str, Any]],
) -> None:
    benchmark = examples["benchmark_customer_example"]
    representative = examples["representative_customer_example"]
    identity = representative["demo_identity"]
    common = representative["benchmark_contract"]

    _add_table(
        document,
        ("구분", "내용"),
        (
            ("작성 기준일", "2026-07-20"),
            ("기준 파일", "data/mock/customer_data_examples.json"),
            ("데이터 성격", "실제 개인정보가 아닌 합성 목데이터"),
            ("활용 목적", "팀 공유·시연 데이터 구조 확인"),
        ),
        (2100, 7260),
    )
    _add_callout(
        document,
        "핵심",
        "대표 시나리오 고객 6명은 1만 명 고객과 동일한 고객·계좌·자산 필드를 모두 "
        "보유하고, 시연용 로그인·시나리오 연결 정보와 실제 적격 ETF 종목 상세가 "
        "추가됩니다.",
    )

    _add_heading(document, "데이터 계약", 1)
    contract_counts = examples["contract_counts"]
    _add_table(
        document,
        ("데이터 층", "필드 수", "저장 내용", "적용 범위"),
        (
            (
                "고객",
                contract_counts["customer_columns"],
                "연령·소득·세액공제·은퇴·성향",
                "1만 명 + 대표 6명",
            ),
            (
                "계좌",
                contract_counts["account_columns_excluding_nested_holdings"],
                "잔액·납입·세액공제·위험비중·수익률",
                "1만 명 + 대표 6명",
            ),
            (
                "자산군 보유",
                contract_counts["benchmark_holding_columns"],
                "자산군·비중·금액·출처",
                "1만 명 + 대표 6명",
            ),
            (
                "ETF 상세 보유",
                contract_counts["detailed_etf_holding_columns"],
                "상품명·종목코드·위험 처리·금액",
                "대표 6명 추가",
            ),
        ),
        (1550, 1000, 3900, 2910),
        numeric_columns={1},
        font_size=8.1,
    )

    benchmark_customer = benchmark["customer"]
    rep_customer = common["customer"]
    _add_heading(document, "대표 예시 한눈에 비교", 1)
    _add_table(
        document,
        ("항목", "일반 고객", "대표 시나리오 고객"),
        (
            ("고객", benchmark_customer["user_id"], identity["nickname"]),
            (
                "연령",
                f"{benchmark_customer['age']}세",
                f"{identity['representative_age']}세",
            ),
            (
                "총급여",
                _money(benchmark_customer["gross_salary_krw"]),
                _money(rep_customer["gross_salary_krw"]),
            ),
            (
                "연금저축+IRP 납입",
                _money(
                    int(benchmark_customer["pension_savings_contribution_krw"])
                    + int(benchmark_customer["irp_contribution_krw"])
                ),
                _money(
                    int(rep_customer["pension_savings_contribution_krw"])
                    + int(rep_customer["irp_contribution_krw"])
                ),
            ),
            (
                "세액공제율",
                f"{benchmark_customer['pension_tax_credit_rate_pct']}%",
                f"{rep_customer['pension_tax_credit_rate_pct']}%",
            ),
            (
                "예상 세액공제",
                _money(benchmark_customer["estimated_pension_tax_credit_krw"]),
                _money(rep_customer["estimated_pension_tax_credit_krw"]),
            ),
            (
                "계좌·총잔액",
                f"{len(benchmark['accounts'])}개 · "
                f"{_money(_account_total(benchmark['accounts']))}",
                f"{len(common['accounts'])}개 · "
                f"{_money(_account_total(common['accounts']))}",
            ),
            ("상세 수준", "자산군 단위", "실제 ETF 종목 단위 추가"),
        ),
        (2100, 3350, 3910),
        font_size=8.1,
    )

    issuer_set = {
        issuer
        for scenario in scenarios
        for account in scenario["accounts"]
        for holding in account["holdings"]
        if holding.get("etf_isu_code")
        for issuer in [_issuer_name(holding["instrument_name"])]
        if issuer is not None
    }
    _add_callout(
        document,
        "ETF 현실화",
        "대표 6명 전체에 "
        + "·".join(sorted(issuer_set))
        + " 상품을 분산 배치했고, KODEX도 유지했습니다.",
    )


def _add_benchmark_customer(document: Document, benchmark: Mapping[str, Any]) -> None:
    customer = benchmark["customer"]
    accounts = benchmark["accounts"]

    _add_page_break(document)
    _add_heading(document, "1. 1만 명 일반 고객 대표 예시", 1)
    _add_callout(
        document,
        "USR00001",
        "30세 근로자 · DC와 개인 IRP 보유 · 안정형·원리금보장 선호 · "
        "DC 방치형 합성 시나리오",
    )
    _add_heading(document, "고객·세액공제·은퇴 정보", 2)
    _add_table(
        document,
        ("영역", "저장값", "영역", "저장값"),
        (
            (
                "고용",
                customer["employment_type"],
                "총급여",
                _money(customer["gross_salary_krw"]),
            ),
            (
                "공제율",
                f"{customer['pension_tax_credit_rate_pct']}%",
                "공제대상 납입",
                _money(customer["total_tax_credit_eligible_contribution_krw"]),
            ),
            (
                "예상 공제",
                _money(customer["estimated_pension_tax_credit_krw"]),
                "연금개시",
                f"{customer['planned_pension_start_age']}세",
            ),
            (
                "위험성향",
                customer["risk_profile"],
                "운용 선호",
                customer["preferred_management_type"],
            ),
            (
                "준비도",
                customer["investment_readiness"],
                "지급 선호",
                customer["payout_preference"],
            ),
            ("데이터 구분", customer["data_kind"], "RAG 저장", "하지 않음"),
        ),
        (1500, 3180, 1500, 3180),
        font_size=8.3,
    )

    _add_heading(document, "계좌 요약", 2)
    account_rows = []
    for account in accounts:
        account_rows.append(
            (
                account["account_type"],
                account["account_id"],
                _money(account["balance_krw"]),
                _money(account["annual_contribution_krw"]),
                _percent(account["risky_asset_ratio"], ratio=True),
                f"{account['trailing_12m_return_pct']}%",
            )
        )
    _add_table(
        document,
        ("계좌", "계좌 ID", "잔액", "연 납입", "위험자산", "과거 12개월"),
        account_rows,
        (900, 1600, 1900, 1750, 1500, 1710),
        numeric_columns={2, 3, 4, 5},
        font_size=8.1,
    )

    _add_heading(document, "계좌별 자산구성", 2)
    holding_rows = []
    for account in accounts:
        by_class = {
            holding["asset_class"]: int(holding["amount_krw"])
            for holding in account["holdings"]
        }
        holding_rows.append(
            (
                account["account_type"],
                _money(by_class.get("EQUITY_KR", 0)),
                _money(by_class.get("EQUITY_GLOBAL", 0)),
                _money(by_class.get("BOND", 0)),
                _money(by_class.get("PRINCIPAL_GUARANTEED", 0)),
                _money(by_class.get("CASH", 0)),
            )
        )
    _add_table(
        document,
        ("계좌", "국내주식", "글로벌주식", "채권", "원리금보장", "현금"),
        holding_rows,
        (850, 1650, 1720, 1500, 1880, 1760),
        numeric_columns={1, 2, 3, 4, 5},
        font_size=7.8,
    )
    _add_body(
        document,
        "해석: 총잔액은 249,250,000원이며, ETF 상품명 대신 국내·글로벌 주식, 채권, "
        "원리금보장, 현금의 자산군 수준으로 저장됩니다.",
        bold_lead="해석:",
    )


def _product_lines(account: Mapping[str, Any]) -> str:
    products = []
    for holding in account["holdings"]:
        code = holding.get("etf_isu_code")
        if code:
            products.append(f"{holding['instrument_name']}({code})")
    return "\n".join(products)


def _add_representative_customer(
    document: Document,
    representative: Mapping[str, Any],
) -> None:
    identity = representative["demo_identity"]
    common = representative["benchmark_contract"]
    customer = common["customer"]
    detailed = representative["detailed_etf_portfolio"]

    _add_page_break(document)
    _add_heading(document, "2. 대표 시나리오 고객 대표 예시", 1)
    _add_callout(
        document,
        identity["nickname"],
        f"기준 고객 {identity['benchmark_user_id']} · "
        f"{identity['representative_age']}세 · "
        "DC·IRP·연금저축의 글로벌 주식형 자산 중복 및 위험 편중 시나리오",
    )

    _add_heading(document, "시연·세액공제 핵심", 2)
    _add_table(
        document,
        ("항목", "저장값", "검증"),
        (
            ("시나리오", identity["scenario_code"], "시연 로그인 연결"),
            ("총급여", _money(customer["gross_salary_krw"]), "5,500만 원 초과"),
            (
                "연금저축 납입",
                _money(customer["pension_savings_contribution_krw"]),
                "600만 원 이하",
            ),
            ("개인 IRP 납입", _money(customer["irp_contribution_krw"]), "생성값"),
            (
                "납입 합산",
                _money(
                    int(customer["pension_savings_contribution_krw"])
                    + int(customer["irp_contribution_krw"])
                ),
                "1,800만 원 이하",
            ),
            (
                "세액공제 대상",
                _money(customer["total_tax_credit_eligible_contribution_krw"]),
                "900만 원 이하",
            ),
            (
                "공제율·예상액",
                f"{customer['pension_tax_credit_rate_pct']}% · "
                f"{_money(customer['estimated_pension_tax_credit_krw'])}",
                "8,760,000 × 13.2% 일치",
            ),
        ),
        (2450, 3900, 3010),
        font_size=8.1,
    )

    _add_heading(document, "계좌·상세 ETF", 2)
    account_by_type = {
        account["account_type"]: account for account in common["accounts"]
    }
    detailed_type_map = {
        "dc": "DC",
        "irp": "IRP",
        "pension_savings": "PENSION_SAVINGS_FUND",
    }
    rows = []
    issuer_set: set[str] = set()
    for account in detailed["accounts"]:
        common_account = account_by_type[detailed_type_map[account["account_type"]]]
        for holding in account["holdings"]:
            if holding.get("etf_isu_code"):
                issuer = _issuer_name(holding["instrument_name"])
                if issuer:
                    issuer_set.add(issuer)
        rows.append(
            (
                common_account["account_type"],
                _money(common_account["balance_krw"]),
                _percent(common_account["risky_asset_ratio"], ratio=True),
                _product_lines(account),
            )
        )
    _add_table(
        document,
        ("계좌", "잔액", "위험자산", "연결된 ETF 상품"),
        rows,
        (900, 1750, 1400, 5310),
        numeric_columns={1, 2},
        font_size=7.3,
    )
    _add_callout(
        document,
        "포트폴리오",
        "이 고객은 "
        + "·".join(sorted(issuer_set))
        + " 등 5개 운용사 ETF 15종을 보유하며, "
        "DC·IRP 일반 위험자산은 각각 70% 이하입니다. 연금저축에는 동일한 위험자산 총량 "
        "한도를 적용하지 않습니다.",
    )
    _add_body(
        document,
        "사용 유의: 모든 고객·계좌 정보는 합성 목데이터입니다. 세액공제액은 교육용 "
        "계산값이며, 수익률은 과거 실적을 참고한 합성값으로 미래 성과를 예측하거나 "
        "보장하지 않습니다.",
        bold_lead="사용 유의:",
    )


def build_document() -> Document:
    examples = json.loads(
        (MOCK_DIR / "customer_data_examples.json").read_text(encoding="utf-8")
    )
    scenarios = json.loads(
        (MOCK_DIR / "chatbot_scenarios.json").read_text(encoding="utf-8")
    )

    document = Document()
    _configure_styles(document)
    _configure_summary_page(document)
    document.core_properties.title = "고객 목데이터 대표 2명 요약"
    document.core_properties.subject = "1만 명 일반 고객과 대표 시나리오 고객 요약"
    document.core_properties.author = "KDA Securities"
    document.core_properties.keywords = "MOCK, 고객데이터, 연금계좌, ETF, 요약"

    _add_title(document)
    _add_overview(document, examples, scenarios)
    _add_benchmark_customer(document, examples["benchmark_customer_example"])
    _add_representative_customer(document, examples["representative_customer_example"])
    return document


def main() -> None:
    document = build_document()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(f"wrote {OUTPUT_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
